from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import random
from generators import generate_formula_image_with_check
from generators import generate_realistic_data, generate_formula_image, generate_chart, fake
from styles import apply_paragraph_format, apply_run_format, apply_title_format
from utils import roman_number, add_page_number, insert_footnote, to_superscript
from config import block_frequencies
from config import margin_ranges
from docx.shared import Length
from PIL import Image 
def create_document(index):
    doc = Document()

    # Устанавливаем отступы для первой секции
    section = doc.sections[0]
    set_random_margins(section)

    # Настройка ориентации страницы
    if random.choice([True, False]):
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT

    base_font_size = random.choice(range(8, 17))
    blocks = []
    footnotes = []

    # Первая секция — одноколоночная (по умолчанию)
    # Добавляем заголовок, который занимает всю ширину страницы
    add_title(doc, base_font_size, num=1)
    blocks.append("title")

    # Добавление заголовка уровня 2 (опционально)
    if random.choice([True, False]):
        add_title(doc, base_font_size, num=2)
        blocks.append("title_level_2")

    # Добавление header и footer
    if random.choice([True, False]):
        add_header(doc, base_font_size)
        blocks.append("header")

    if random.choice([True, False]):
        add_footer(doc)
        blocks.append("footer")

    # Добавляем секцию с многоколоночным форматом
    is_multicolumn = random.choices([True, False], weights=[75, 25], k=1)[0]  # Изменены веса
    if is_multicolumn:
        num_columns = random.choice([2, 3])
        new_section = doc.add_section(WD_SECTION.CONTINUOUS)
        set_random_margins(new_section)
        column_width, space = add_multicolumn_section(new_section, num_columns=num_columns)
    else:
        num_columns = 1
        column_width = page_width = section.page_width.twips
    print(num_columns)
    # Добавляем основные блоки
    num_blocks = random.randint(20, 40)
    block_types = list(block_frequencies.keys())
    weights = list(block_frequencies.values())

    for _ in range(num_blocks):
        block = random.choices(block_types, weights=weights, k=1)[0]

        if block == 'paragraph':
            add_paragraph(doc, base_font_size, footnotes, column_width)
        elif block == 'table':
            add_table(doc, base_font_size, column_width)
        elif block == 'picture':
            add_picture(doc, base_font_size, column_width)
        elif block == 'numbered_list':
            add_numbered_list(doc, base_font_size, column_width)
        elif block == 'marked_list':
            add_marked_list(doc, base_font_size, column_width)
        elif block == 'formula':
            add_formula(doc, base_font_size, column_width)
        elif block == 'title':
            add_title(doc, base_font_size, num=2)
        blocks.append(block)

    add_footnotes_section(doc, footnotes, base_font_size)

    # Сохраняем документ
    output_dir = "docx"
    os.makedirs(output_dir, exist_ok=True)

    filename = os.path.join(output_dir, f"doc_{index}.docx")
    doc.save(filename)

    return filename, base_font_size

def add_multicolumn_section(section, num_columns=2):
    """
    Настраивает многоколоночную секцию в документе.
    """
    space = random.randint(400, 700)
    # Получаем доступные размеры страницы
    page_width = section.page_width.twips - (section.left_margin.twips + section.right_margin.twips)

    # Рассчитываем общий отступ между колонками
    total_spacing = (num_columns - 1) * space

    # Проверяем, что ширина колонок не превышает доступную ширину страницы
    if total_spacing >= page_width:
        raise ValueError("Ширина отступов между колонками превышает доступную ширину страницы.")
    
    # Рассчитываем ширину колонок
    column_width = (page_width - total_spacing) // num_columns

    # Настройка колонок
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0] if sectPr.xpath('./w:cols') else OxmlElement('w:cols')
    cols.set(qn('w:num'), str(num_columns))
    cols.set(qn('w:space'), str(space))
    sectPr.append(cols)

    return column_width, space



def add_full_width_block(doc, add_function, base_font_size, num=1):
    """
    Добавляет блок, который занимает всю ширину страницы (оборачивая в отдельную секцию с одной колонкой).

    :param doc: объект Document
    :param add_function: функция добавления блока (например, add_title)
    :param base_font_size: базовый размер шрифта
    :param num: номер заголовка (для различных стилей)
    """
    # Добавляем новую секцию с одной колонкой без разрыва страницы
    new_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_random_margins(new_section)
    add_multicolumn_section(new_section, num_columns=1)

    # Добавляем блок
    add_function(doc, base_font_size, num=num)

    # Возвращаемся к предыдущей секции (многоколоночной)
    # Добавляем новую секцию с многоколоночным форматом
    multi_column_section = doc.add_section(WD_SECTION.CONTINUOUS)
    num_columns = random.choice([2, 3])
    add_multicolumn_section(multi_column_section, num_columns=num_columns)
    column_width, space = add_multicolumn_section(multi_column_section, num_columns=num_columns)

    return column_width

def remove_table_borders(table):
    """
    Функция скрывает границы ячеек таблицы.
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')

    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        borders.append(border)

    tblPr.append(borders)

def add_paragraph(doc, base_font_size, footnotes, column_width):
    """
    Добавляет параграф в документ.

    :param doc: объект Document из python-docx
    :param base_font_size: базовый размер шрифта
    :param footnotes: список сносков
    :param column_width: ширина колонки для ограничения ширины параграфа
    """
    paragraph = doc.add_paragraph()

    # Генерация текста абзаца
    text = fake.text(max_nb_chars=random.randint(100, 1200))
    insert_footnote(paragraph, text, base_font_size, footnotes, footnote_per_page=5)
    run = paragraph.add_run(text)
    run.font.size = Pt(base_font_size)
    # Форматирование абзаца
    apply_paragraph_format(paragraph, base_font_size, is_multicolumn=True)

    if paragraph.runs:
        run = paragraph.runs[0]
        apply_run_format(run, base_font_size)

    paragraph.paragraph_format.space_after = Pt(12)

def add_title(doc, base_font_size, num=1):
    """
    Добавляет заголовок в документ.

    :param doc: объект Document из python-docx
    :param base_font_size: базовый размер шрифта
    :param num: номер заголовка (1 для основного, 2 для подзаголовка и т.д.)
    """
    title_number = random.choice([True, False])
    title = doc.add_paragraph()
    run = title.add_run()
    if num == 1:
        title_text = fake.sentence(nb_words=random.randint(3, 10))
    elif num == 2:
        title_text = fake.sentence(nb_words=random.randint(2, 6))
    else:
        title_text = fake.sentence(nb_words=random.randint(1, 5))

    if title_number:
        title_text = f"{random.randint(1, 100)}. {title_text}"

    run.text = title_text

    apply_title_format(title, base_font_size, num=num)

# add_table_caption, remove_table_borders, generate_realistic_data, insert_footnote

def add_table(doc, base_font_size, column_width_twips):
    """
    Добавляет таблицу в документ с различными стилями, границами, цветами и выравниванием.

    :param doc: объект Document из python-docx
    :param base_font_size: базовый размер шрифта
    :param column_width_twips: ширина колонки в twips для ограничения ширины таблицы
    """
    # Перевод twips в дюймы (1 дюйм = 1440 twips)
    column_width_inch = column_width_twips / 1440
    
    # Генерация таблицы
    add_table_caption(doc, base_font_size)
    
    if column_width_inch <= 3:
        num_cols = random.randint(2, 4)
        num_rows = random.randint(2, 5)
    else:
        num_cols = random.randint(3, 6)
        num_rows = random.randint(3, 10)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    # Ограничение ширины таблицы
    table.autofit = False
    table_width = column_width_twips # Устанавливаем ширину таблицы в дюймах
    table.width = Inches(table_width)
    
    # Устанавливаем ширину каждой колонки
    col_width = table_width / num_cols
    for col in table.columns:
        col.width = Inches(col_width)
    
    # Удаление границ таблицы
    if random.choice([True, False]):
        remove_table_borders(table)
    
    # Применение случайного стиля
    table.style = random.choice(['Table Grid', 'Light Shading', 'Light Grid', 'Medium Shading 1', 'Medium List 1'])
    
    # Случайное выравнивание таблицы
    table.alignment = random.choice([
        WD_ALIGN_PARAGRAPH.JUSTIFY
    ])
    
    # Случайное выравнивание текста в ячейках
    cell_alignment = random.choice([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.JUSTIFY
    ])
    
    # Настройка цвета текста и фона строк
    text_color_choice = random.choice(['black', 'white'])
    text_rgb = RGBColor(0, 0, 0) if text_color_choice == 'black' else RGBColor(255, 255, 255)
    
    row_colors = [
        (255, 255, 255), (220, 220, 220), (240, 248, 255), (255, 250, 205), (255, 228, 225)
    ] if text_color_choice == 'black' else [
        (0, 0, 0), (105, 105, 105), (25, 25, 112), (139, 0, 0), (0, 100, 0)
    ]
    
    # Случайные цвета строк
    color1, color2 = random.sample(row_colors, 2) if len(row_colors) >= 2 else (row_colors[0], row_colors[0])
    
    # Заполнение таблицы
    for i, row in enumerate(table.rows):
        bg_color = color1 if i % 2 == 0 else color2  # Чередование цветов строк
        for cell in row.cells:
            # Генерация содержимого ячейки
            text = generate_realistic_data()
            insert_footnote(cell.paragraphs[0], text, base_font_size, [], footnote_per_page=0)
    
            # Добавление текста, если нет run
            if not cell.paragraphs[0].runs:
                run = cell.paragraphs[0].add_run(cell.text)
            else:
                run = cell.paragraphs[0].runs[0]
    
            # Настройка шрифта
            run.font.size = Pt(base_font_size)
            run.font.color.rgb = text_rgb
            run.font.bold = random.choice([True, False])
            run.font.italic = random.choice([True, False])
    
            # Настройка выравнивания текста
            cell.paragraphs[0].alignment = cell_alignment
    
            # Настройка фона ячейки
            tcPr = cell._element.get_or_add_tcPr()
            shd = tcPr.find(qn('w:shd'))
            if shd is None:
                shd = OxmlElement("w:shd")
                tcPr.append(shd)
            shd_color = '{:02X}{:02X}{:02X}'.format(*bg_color)
            shd.set(qn('w:fill'), shd_color)
            shd.set(qn('w:val'), 'clear')
    
    # Добавление абзаца после таблицы
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)
    
    return table


def add_picture(doc, base_font_size, column_width_twips):
    """
    Добавляет изображение или график в документ. Масштабирует изображение до ширины колонки,
    если оно слишком широкое для страницы.

    :param doc: объект Document из python-docx
    :param base_font_size: базовый размер шрифта
    :param column_width_twips: ширина колонки в twips для ограничения ширины изображения
    """
    from PIL import Image
    # Определяем, использовать ли график (30% вероятность)
    use_chart = random.choices([True, False], weights=[0.3, 0.7], k=1)[0]
    location_signature_after = random.choices([True, False], weights=[0.7, 0.3], k=1)[0]

    # Выбор изображения или генерация графика
    if use_chart:
        image_path = generate_chart()
    else:
        image_folder = "document_generator/Dataset_images"  # Директория с изображениями
        if os.path.isdir(image_folder):
            if not hasattr(add_picture, "cached_images"):
                add_picture.cached_images = [f for f in os.listdir(image_folder) if f.lower().endswith(('.png', '.jpg', '.jpeg'))]

            if add_picture.cached_images:
                image_path = os.path.join(image_folder, random.choice(add_picture.cached_images))
            else:
                print("No images found in folder.")
                return
        else:
            print("Image folder does not exist.")
            return

    # Создание параграфа для изображения
    paragraph = doc.add_paragraph()
    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]

    # Добавление подписи перед изображением, если нужно
    if not location_signature_after:
        add_picture_caption(doc, base_font_size)

    # Вставка изображения в документ с автоматическим масштабированием
    try:
        # Рассчитываем максимальную ширину изображения в дюймах
        max_img_width = column_width_twips  # Максимальная ширина изображения	
        width_img = int(max_img_width * random.uniform(0.4, 0.8)) / 1440
        # Вставка изображения в документ
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(width_img))  

    except Exception as e:
        print(f"Ошибка при вставке изображения из {image_path}: {e}")

    # Добавление подписи после изображения, если нужно
    if location_signature_after:
        add_picture_caption(doc, base_font_size)

    # Удаление временного файла с графиком или масштабированного изображения
    if use_chart and os.path.exists(image_path):
        os.remove(image_path)


def add_picture_caption(doc, base_font_size):
    """
    Добавляет подпись к изображению.

    :param doc: объект Document из python-docx
    :param base_font_size: базовый размер шрифта
    """
    img_text = random.choice(["Рис. ", "Рисунок "])
    number = str(random.randint(1, 100)) + (random.choice(["-", ""]) if img_text == "Рисунок " else "")
    caption_text = img_text + number

    paragraph = doc.add_paragraph()
    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]

    run = paragraph.add_run(caption_text)
    run.font.size = Pt(base_font_size - random.choice([1, 2, 0, 3]))
    run.font.italic = True

def add_table_caption(doc, base_font_size):
    """
    Добавляет подпись к таблице.

    :param doc: объект Document из python-docx
    :param base_font_size: базовый размер шрифта
    """
    tbl_text = random.choice(["Табл. ", "Таблица "])
    number = str(random.randint(1, 100)) + (random.choice(["-", ""]) if tbl_text == "Таблица " else "")
    caption_text = tbl_text + number

    paragraph = doc.add_paragraph()
    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]

    run = paragraph.add_run(" " + caption_text)
    run.font.size = Pt(base_font_size - random.choice([1, 2, 0, 3]))
    run.font.italic = True

def add_numbered_list(doc, base_font_size, column_width):
    """
    Добавляет нумерованный список в документ.

    :param doc: объект Document из python-docx
    :param base_font_size: базовый размер шрифта
    :param column_width: ширина колонки для ограничения выравнивания
    """
    numbering_styles = ['1.', 'I.', 'A.', 'a.', 'i.', '1)', 'I)', 'A)', 'a)', 'i)']
    style_choice = random.choice(numbering_styles) 

    line_spacing = 0.8
    left_indent = Pt(random.randint(5, 15))
    count = random.randint(2, 7)

    for i in range(1, count + 1):
        paragraph = doc.add_paragraph()

        if style_choice in ['1.', '1)']:
            number = f"{i}{style_choice[-1]}"
        elif style_choice in ['I.', 'I)']:
            number = f"{roman_number(i)}{style_choice[-1]}"
        elif style_choice in ['A.', 'A)']:
            number = f"{chr(64 + i)}{style_choice[-1]}"
        elif style_choice in ['a.', 'a)']:
            number = f"{chr(96 + i)}{style_choice[-1]}"
        elif style_choice in ['i.', 'i)']:
            number = f"{roman_number(i).lower()}{style_choice[-1]}"

        run = paragraph.add_run(f"{number} {fake.sentence(nb_words=random.randint(3, 10))}")
        run.font.size = Pt(base_font_size)

        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.left_indent = left_indent
    paragraph.paragraph_format.space_after = Pt(12)    

def add_marked_list(doc, base_font_size, column_width):
    """
    Добавляет маркированный список в документ.

    :param doc: объект Document из python-docx
    :param base_font_size: базовый размер шрифта
    :param column_width: ширина колонки для ограничения выравнивания
    """
    markers = ['•', '○', '■', '–', '*', '→']
    marker = random.choice(markers)  

    line_spacing = 0.8
    left_indent = Pt(random.randint(5, 15))
    count = random.randint(2, 7)

    for _ in range(count):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.line_spacing = line_spacing  
        paragraph.paragraph_format.left_indent = left_indent  

        text_with_marker = f"{marker} {fake.sentence(nb_words=random.randint(3, 10))}"
        run = paragraph.add_run(text_with_marker)
        run.font.size = Pt(base_font_size)
        
    paragraph.paragraph_format.space_after = Pt(12)

def add_header(doc, base_font_size):
    """
    Добавляет заголовок (header) в документ.
    """
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = fake.text(max_nb_chars=random.randint(5, 20))
    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]
    run = paragraph.runs[0]
    run.font.size = Pt(base_font_size - 1)
    run.font.italic = True

def add_footer(doc):
    """
    Добавляет нижний колонтитул (footer) в документ.
    """
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]
    page_number = random.choice([True, False])
    if page_number:
        add_page_number(paragraph)
    else:
        paragraph.add_run(fake.word())

def add_footnotes_section(doc, footnotes, base_font_size):
    """
    Добавляет раздел сносков в конец документа.
    """
    if not footnotes:
        return

    doc.add_paragraph("_" * 50)

    for number, text in footnotes:
        footnote_paragraph = doc.add_paragraph()
        footnote_run = footnote_paragraph.add_run(f"{to_superscript(number)}. {text}")
        footnote_run.font.size = Pt(base_font_size)
        footnote_paragraph.paragraph_format.space_before = Pt(2)

def add_formula(doc, base_font_size, column_width_twips):
    """
    Добавляет формулу в документ как изображение.
    """
    image_stream = generate_formula_image_with_check()
    column_width_in = column_width_twips / 1440  # Twips → Inches

    paragraph = doc.add_paragraph()

    # Ограничиваем ширину формулы
    if column_width_in < 2:
        max_width = column_width_in
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif column_width_in <= 3.1:
        max_width = 2  # Задаём фиксированную ширину
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        max_width = 3.1
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run()
    try:
        run.add_picture(image_stream, width=Inches(max_width))
    except Exception as e:
        print(f"Ошибка при вставке формулы: {e}")
def set_random_margins(section):
    """
    Устанавливает случайные отступы для секции документа.
    Случайно выбирает, будут ли левая и правая границы равны.
    """
    # Случайный выбор равенства отступов
    equal_margins = random.choice([True, False])

    if equal_margins:
        # Устанавливаем равные отступы
        side_margin = random.uniform(*margin_ranges['left'])
        section.left_margin = Inches(side_margin)
        section.right_margin = Inches(side_margin)
    else:
        # Устанавливаем разные отступы
        left_margin = random.uniform(*margin_ranges['left'])
        right_margin = random.uniform(*margin_ranges['right'])
        section.left_margin = Inches(left_margin)
        section.right_margin = Inches(right_margin)

    # Устанавливаем верхний и нижний отступы
    section.top_margin = Inches(random.uniform(*margin_ranges['top']))
    section.bottom_margin = Inches(random.uniform(*margin_ranges['bottom']))