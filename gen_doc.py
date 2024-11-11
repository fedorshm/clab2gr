import os
import random
from collections import OrderedDict
import csv  
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH  
from docx.enum.table import WD_TABLE_ALIGNMENT  
from docx.enum.section import WD_ORIENTATION  
from docx.oxml.ns import qn  
from docx.oxml import OxmlElement  

from faker import Faker  
from mimesis import Person, Datetime  
from mimesis.locales import Locale  

import sympy as sp
import matplotlib.pyplot as plt

from io import BytesIO

locales = OrderedDict([('en-US', 1), ('ru-RU', 2)])
fake = Faker(locales)

def create_document(index):
    doc = Document()

    section = doc.sections[0]
    if random.choice([True, False]):
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width

    base_font_size = random.choice(range(8, 17))

    if random.choice([True, False]):
        add_title(doc, base_font_size)

    if random.choice([True, False]):
        add_paragraph(doc, base_font_size)
    if random.choice([True, False]):    
        add_header(doc, base_font_size)
    if random.choice([True, False]):      
        add_footer(doc)

    if random.choice([True, False]):  
        add_table(doc, base_font_size)

    if random.choice([True, False]):
        add_paragraph(doc, base_font_size)

    if random.choice([True, False]):  
        add_picture(doc, base_font_size)
    if random.choice([True, False]):      
        add_numbered_list(doc, base_font_size)
    if random.choice([True, False]):      
        add_marked_list(doc, base_font_size)
    if random.choice([True, False]):      
        add_formula(doc, base_font_size) 

    if random.choice([True, False]):  
        add_table(doc, base_font_size)  

    #if random.choice([True, False]):
    #    add_footnote(doc, base_font_size)

    output_dir = "docx"
    os.makedirs(output_dir, exist_ok=True)  

    filename = os.path.join(output_dir, f"doc_{index}.docx")
    doc.save(filename)
    
    return filename, base_font_size  
def create_annotation():
    pass

def generate_formula_image():
    symbols_dict = {
        'variables': sp.symbols('x y z a b c 1 2 3 4 5 6 7 8 9'),
        'greek': sp.symbols('alpha beta gamma delta epsilon theta lambda mu nu sigma phi psi omega chi'),
        'constants': [sp.pi, sp.E],
        'functions': [sp.sin, sp.cos, sp.exp, sp.log, sp.sqrt, sp.tan, sp.cot, sp.sec, sp.csc, sp.sinh, sp.cosh, sp.tanh],
    }

    def random_symbol():
        category = random.choice(['variables', 'greek', 'constants'])
        return random.choice(symbols_dict[category])

    def random_expression(depth=0):
        if depth > 2 or random.choice([True, False]):
            expr = random_symbol()
            if random.choice([True, False]):
                func = random.choice(symbols_dict['functions'])
                expr = func(expr)
            return expr
        else:
            left = random_expression(depth + 1)
            right = random_expression(depth + 1)
            operator = random.choice(['+', "-", '-', '*', '/', '**'])
            if operator == '+':
                return left + right
            elif operator == '-':
                return left - right
            elif operator == '*':
                return left * right
            elif operator == '/':
                right = right + sp.Integer(1)
                return left / right
            elif operator == '**':
                exponent = random.randint(2, 3)
                return left ** exponent
            else:
                return left

    expr = random_expression()
    latex_formula = sp.latex(expr)

    plt.figure(figsize=(5, 1.5))
    plt.text(0.5, 0.5, f"${latex_formula}$", fontsize=20, ha='center', va='center')
    plt.axis('off')

    image_stream = BytesIO()
    plt.savefig(image_stream, format='png', bbox_inches='tight', pad_inches=0.1, transparent=True)
    plt.close()
    image_stream.seek(0) 

    return image_stream

def remove_table_borders(table):
    tbl = table._tbl  
    tblPr = tbl.tblPr  
    borders = OxmlElement('w:tblBorders')

    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil') 
        borders.append(border)

    tblPr.append(borders)

def add_paragraph(doc, base_font_size):
    col_paragraph = random.choices([True, False], weights=[0.2, 0.8], k=1)[0]

    if col_paragraph:
        num_columns = random.choice([2, 3])  
        table = doc.add_table(rows=1, cols=num_columns)  
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        remove_table_borders(table)

        for cell in table.rows[0].cells:
            paragraph = cell.paragraphs[0]
            paragraph.text = fake.text(max_nb_chars=random.randint(100, 500))
            paragraph_format = paragraph.paragraph_format

            paragraph_format.first_line_indent = Pt(35.5) if random.choice([True, False]) else None

            paragraph_format.alignment = random.choices([
                WD_ALIGN_PARAGRAPH.JUSTIFY,
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.RIGHT
            ], weights=[0.5, 0.25, 0.25], k=1)[0]

            run = paragraph.runs[0]
            run.font.size = Pt(base_font_size)
            run.font.name = random.choice(['Times New Roman', 'Arial', 'Calibri', 'Verdana'])
    else:
        paragraph = doc.add_paragraph(fake.text(max_nb_chars=random.randint(50, 1000)))
        paragraph_format = paragraph.paragraph_format

        paragraph_format.first_line_indent = Pt(35.5) if random.choice([True, False]) else None

        paragraph_format.alignment = random.choices([
            WD_ALIGN_PARAGRAPH.JUSTIFY,
            WD_ALIGN_PARAGRAPH.CENTER,
            WD_ALIGN_PARAGRAPH.LEFT,
            WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.3, 0.3, 0.2, 0.2], k=1)[0]

        run = paragraph.runs[0]
        run.font.size = Pt(base_font_size)
        run.font.name = random.choice(['Times New Roman', 'Arial', 'Calibri', 'Verdana'])

def add_title(doc, base_font_size):
    title_number = random.choice([True, False])
    title = doc.add_paragraph()
    run = title.add_run()
    
    title_text = fake.sentence(nb_words=random.randint(1, 50))    
    if title_number:
        title_text = f"{random.randint(0, 100)}. {title_text}"
    
    run.text = title_text
    
    run.font.bold = random.choice([True, False])
    run.font.italic = random.choice([True, False])
    
    title_font_size = base_font_size + random.choice([2, 3, 4, 5])
    run.font.size = Pt(title_font_size)
    
    title_alignment = random.choices(
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        weights=[0.6, 0.2, 0.2],
        k=1
    )[0]
    title.alignment = title_alignment

def add_table(doc, base_font_size):
    add_table_caption(doc, base_font_size)
    table = doc.add_table(rows=random.randint(3, 10), cols=random.randint(3, 7))
    
    remove_border = random.choice([True, False])
    if remove_border:
        remove_table_borders(table)
    
    side_borders = False
    if remove_border:
        side_borders = random.choice([True, False])
        if side_borders:
            tbl = table._tbl
            tblPr = tbl.tblPr
            borders = OxmlElement('w:tblBorders')

            for border_name in ('left', 'right', 'insideH', 'insideV'):
                border_element = OxmlElement(f'w:{border_name}')
                border_element.set(qn('w:val'), 'nil')
                borders.append(border_element)
            tblPr.append(borders)
    
    apply_style = random.choice(['Table Grid'])
    table.style = apply_style
   
    table.alignment = random.choice([
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
    ])
    
    cell_alignment = random.choice([
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
    ])

    text_color_choice = random.choice(['black', 'white'])
    if text_color_choice == 'black':
        text_rgb = RGBColor(0, 0, 0)
        
        row_colors = [
            (255, 255, 255),  
            (220, 220, 220),  
            (240, 248, 255),  
            (255, 250, 205),  
            (255, 228, 225),  
        ]
    else:
        text_rgb = RGBColor(255, 255, 255)
        row_colors = [
            (0, 0, 0),        
            (105, 105, 105),  
            (25, 25, 112),    
            (139, 0, 0),      
            (0, 100, 0),      
        ]

    if len(row_colors) >= 2:
        color1, color2 = random.sample(row_colors, 2)
    else:
        color1, color2 = row_colors[0], row_colors[0]

    for i, row in enumerate(table.rows):
        bg_color = color1 if i % 2 == 0 else color2
        for cell in row.cells:
            cell.text = fake.text(max_nb_chars=random.randint(5, 100))
            
            if not cell.paragraphs[0].runs:
                run = cell.paragraphs[0].add_run(cell.text)
            else:
                run = cell.paragraphs[0].runs[0]
            
            run.font.size = Pt(base_font_size)
            cell.paragraphs[0].alignment = cell_alignment
            
            tcPr = cell._element.get_or_add_tcPr()
            shd = tcPr.find(qn('w:shd'))
            if shd is None:
                shd = OxmlElement("w:shd")
                tcPr.append(shd)
            shd_color = '{:02X}{:02X}{:02X}'.format(*bg_color)
            shd.set(qn('w:fill'), shd_color)
            shd.set(qn('w:val'), 'clear')  
            
            run.font.color.rgb = text_rgb

def add_picture(doc, base_font_size):
    location_signature_after = random.choices(
        [True, False],
        weights=[0.7, 0.3],
        k=1
    )[0]
    image_folder = "Dataset_images"
    
    if os.path.isdir(image_folder):
        images = [f for f in os.listdir(image_folder) if f.endswith(('.png', '.jpg', '.jpeg'))]
        if images:
            image_path = os.path.join(image_folder, random.choice(images))
            
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            
            if not location_signature_after:
                add_picture_caption(doc, base_font_size)
            
            run.add_picture(image_path)
            
            if location_signature_after:
                add_picture_caption(doc, base_font_size)
        else:
            print("No images found in folder.")
    else:
        print("Folder is empty.")
    

def add_picture_caption(doc, base_font_size):
    img_text = random.choice(["Рис. ", "Рисунок "])
    number = str(random.randint(1, 100)) + (random.choice(["-", ""]) if img_text == "Рисунок " else "")
    caption_text = img_text + number
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]
    
    run = paragraph.add_run()
    run.add_text(" " + caption_text)  
    
    run.font.size = Pt(base_font_size - random.choice([1, 2, 0, 3]))
    run.font.italic = True  

def add_table_caption(doc, base_font_size):
    tbl_text = random.choice(["Табл. ", "Таблица "])
    number = str(random.randint(1, 100)) + (random.choice(["-", ""]) if tbl_text == "Таблица " else "")
    caption_text = tbl_text + number
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]
    run = paragraph.add_run()
    run.add_text(" " + caption_text)  
    
    run.font.size = Pt(base_font_size - random.choice([1, 2, 0, 3]))
    run.font.italic = True  

def add_numbered_list(doc, base_font_size):
    count = random.randint(2,7)
    for _ in range(count):
        paragraph = doc.add_paragraph(style='List Number')
        run = paragraph.add_run(fake.sentence(nb_words=random.randint(3, 10)))
        run.font.size = Pt(base_font_size)
        paragraph.paragraph_format.line_spacing = 1.0

def add_marked_list(doc, base_font_size):
    count = random.randint(2,7)
    for _ in range(count):
        paragraph = doc.add_paragraph(style='List Bullet')
        run = paragraph.add_run(fake.sentence(nb_words=random.randint(3, 10)))
        run.font.size=Pt(base_font_size)
        paragraph.paragraph_format.line_spacing = 1.0

def add_header(doc, base_font_size):
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = fake.text(max_nb_chars=random.randint(5, 20))
    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]
    run = paragraph.runs[0]
    run.font.size = Pt(base_font_size - 1)
    run.font.italic = True

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]
    page_number = random.choice([True, False])
    if page_number:
        add_page_number(paragraph)
    else:
        paragraph.add_run(fake.word())     

def add_page_number(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def to_superscript(number):
    """
    Преобразует число в его суперскриптный аналог.
    Например, 1 -> ¹, 2 -> ² и т.д.
    """
    superscript_map = str.maketrans('0123456789', '⁰¹²³⁴⁵⁶⁷⁸⁹')
    return str(number).translate(superscript_map)

def insert_footnote(paragraph, text, base_font_size, current_page_footnotes, footnote_per_page):
    """
    Вставляет сноску в текст абзаца с вероятностью 20% и ограничением по количеству сносок на страницу.
    
    :param paragraph: Абзац, в который вставляется текст.
    :param text: Текст абзаца.
    :param base_font_size: Базовый размер шрифта.
    :param current_page_footnotes: Список сносок текущей страницы.
    :param footnote_per_page: Максимальное количество сносок на страницу.
    :return: None
    """
    words = text.split()
    if len(words) < 3:
        run = paragraph.add_run(text)
        run.font.size = Pt(base_font_size)
        return

    # Решаем, вставлять ли сноску (например, с вероятностью 20%)
    if random.random() < 0.2 and len(current_page_footnotes) < footnote_per_page:
        footnote_number = len(current_page_footnotes) + 1
        footnote_number_sup = to_superscript(footnote_number)

        # Выбираем случайное место для вставки сноски
        insert_pos = random.randint(1, len(words)-1)
        new_text = ' '.join(words[:insert_pos]) + f' {footnote_number_sup} ' + ' '.join(words[insert_pos:])
        run = paragraph.add_run(new_text)
        run.font.size = Pt(base_font_size)
        
        footnote_text = fake.text(max_nb_chars=random.randint(20, 50))
        current_page_footnotes.append((footnote_number, footnote_text))
    else:
        run = paragraph.add_run(text)
        run.font.size = Pt(base_font_size)

def add_footnotes_section(doc, footnotes, base_font_size):
    """
    Добавляет раздел с сносками в конец страницы.
    
    :param doc: Объект Document.
    :param footnotes: Список кортежей (номер сноски, текст сноски).
    """
    if not footnotes:
        return

    doc.add_paragraph("_" * 50)

    # Добавляем каждую сноску
    for number, text in footnotes:
        footnote_paragraph = doc.add_paragraph()
        footnote_run = footnote_paragraph.add_run(f"{number}. {text}")
        footnote_run.font.size = Pt(base_font_size)
        footnote_paragraph.paragraph_format.space_before = Pt(2)

def add_formula(doc, base_font_size):
    image_stream = generate_formula_image()

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  
    run = paragraph.add_run()
    try:
        run.add_picture(image_stream, width=Pt(300))  
    except Exception as e:
        print(f"Ошибка при вставке формулы: {e}")

def main():
    all_data = []  # Список для хранения данных о документах

    for index in range(1000):  
        filename, base_font_size = create_document(index)
        base_filename = os.path.splitext(os.path.basename(filename))[0]  # Получаем имя без расширения
        all_data.append({'filename': base_filename, 'base_font_size': base_font_size})
        print(f"Сгенерирован документ: {filename} с базовым размером шрифта: {base_font_size}")

    # Запись данных в CSV
    output_csv = 'font_sizes.csv'
    with open(output_csv, 'w', newline='', encoding='utf-8') as csvfile:
        fieldnames = ['filename', 'base_font_size']
        writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
        writer.writeheader()
        for row in all_data:
            writer.writerow(row)

    print(f"Данные о базовых размерах шрифтов сохранены в {output_csv}")

if __name__ == "__main__":
    main()
