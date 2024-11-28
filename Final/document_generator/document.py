# document.py

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import random

from docx.shared import Inches

from generators import generate_realistic_data, generate_formula_image, fake
from styles import apply_paragraph_format, apply_run_format, apply_title_format
from utils import roman_number, add_page_number, insert_footnote, to_superscript
from config import block_frequencies, MARGIN_RANGES

def create_document(index):
    doc = Document()

    section = doc.sections[0]
    
    set_random_margins(doc)

    if random.choice([True, False]):
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    else:
        section.orientation = WD_ORIENTATION.PORTRAIT
    print(section.page_width.twips)
    is_multicolumn = random.choices([True, False], weights=[25, 75], k=1)[0]

    base_font_size = random.choice(range(8, 17))
    blocks = []
    footnotes = []

    add_title(doc, base_font_size,1)
    blocks.append("title")

    add_header(doc, base_font_size)
    blocks.append("header")

    add_footer(doc)
    blocks.append("footer")

    num_blocks = random.randint(10, 20)

    block_types = list(block_frequencies.keys())
    weights = list(block_frequencies.values())

    current_column = 0  

    if is_multicolumn:
        main_table, num_columns, column_width = add_multicolumn_table(doc)
    else:
        main_table, num_columns, column_width = None, 1, section.page_width.twips   

    for _ in range(num_blocks):
        block = random.choices(block_types, weights=weights, k=1)[0]
        if is_multicolumn:
            cell = main_table.rows[0].cells[current_column]
        else:
            cell = None  
        
            add_paragraph(doc, base_font_size, footnotes, cell)
        if block == 'paragraph':
            add_paragraph(doc, base_font_size, footnotes, cell)
        elif block == 'table':
            add_table(doc, base_font_size, column_width, cell)
            if cell is not None:
                paragraph = cell.add_paragraph()
            else:
                paragraph = doc.add_paragraph()
        elif block == 'picture':
            add_picture(doc, base_font_size, cell)
        elif block == 'numbered_list':
            add_numbered_list(doc, base_font_size, cell)
        elif block == 'marked_list':
            add_marked_list(doc, base_font_size, cell)
        elif block == 'formula':
            add_formula(doc, base_font_size, cell)
        blocks.append(block)

        if is_multicolumn:
            current_column = (current_column + 1) % num_columns

    add_footnotes_section(doc, footnotes, base_font_size)

    output_dir = "docx"
    os.makedirs(output_dir, exist_ok=True)

    filename = os.path.join(output_dir, f"doc_{index}.docx")
    doc.save(filename)

    return filename, base_font_size

def set_random_margins(doc):
    """
    Устанавливает случайные отступы на основе диапазонов из конфигурации.
    """
    left_margin = random.uniform(MARGIN_RANGES['left'][0], MARGIN_RANGES['left'][1])
    right_margin = left_margin if random.random() > 0.2 else random.uniform(MARGIN_RANGES['right'][0], MARGIN_RANGES['right'][1])  # 80% вероятность, что правый и левый отступы будут одинаковыми
    top_margin = random.uniform(MARGIN_RANGES['top'][0], MARGIN_RANGES['top'][1])
    bottom_margin = random.uniform(MARGIN_RANGES['bottom'][0], MARGIN_RANGES['bottom'][1])

    # Применяем отступы
    section = doc.sections[0]
    section.left_margin = Inches(left_margin)
    section.right_margin = Inches(right_margin)
    section.top_margin = Inches(top_margin)
    section.bottom_margin = Inches(bottom_margin)

import random
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_multicolumn_table(doc, min_columns=2, max_columns=3):
    """
    Реализует мультиколоночность документа с случайными отступами от колонок.
    """
    num_columns = random.choice([min_columns, max_columns]) 
    table = doc.add_table(rows=1, cols=num_columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False  

    remove_table_borders(table)  # Убираем границы таблицы

    section = doc.sections[0]
    page_width = section.page_width.twips  
    
    # Рассчитываем отступы
    # Генерация случайных отступов между колонками
    total_spacing = 0
    column_spacing = []

    for i in range(num_columns - 1):  # Отступы между колонками
        # Случайный отступ для каждой пары колонок
        inch_to_twips = 1440  # 1 дюйм = 1440 twips
        spacing_in_twips = random.uniform(0.1, 0.5) * inch_to_twips
        spacing = random.uniform(0.1, 0.5) * inch_to_twips   # от 0.1 до 0.5 дюйма
        column_spacing.append(spacing)
        total_spacing += spacing

    # Рассчитываем ширину колонок с учётом отступов
    column_width = (page_width - total_spacing) / num_columns

    for col in range(num_columns):
        table.columns[col].width = int(column_width) * 1440

    # Настройка отступов внутри ячеек таблицы (если необходимо)
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblCellMar = OxmlElement('w:tblCellMar')
    
    for margin in ['top', 'start', 'bottom', 'end']:
        mar = OxmlElement(f'w:{margin}')
        mar.set(qn('w:w'), '100')  
        mar.set(qn('w:type'), 'dxa')
        tblCellMar.append(mar)
    
    tblPr.append(tblCellMar)

    return table, num_columns, column_width  


def remove_table_borders(table):
    """
    Функция скрывает границы ячеек таблицы"
    """
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')

    for border_name in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        borders.append(border)

    tblPr.append(borders)

def add_paragraph(doc, base_font_size, footnotes, cell=None):
    """
    Добавляет параграф в документ или в указанную ячейку таблицы.
    """
    if random.choices([True,False],[0.1,0.9], k=1)[0]:
        add_title(doc, base_font_size, 2)
    if cell is not None:
        paragraph = cell.add_paragraph()
    else:
        paragraph = doc.add_paragraph()

    # Генерация текста абзаца
    text = fake.text(max_nb_chars=random.randint(50, 1000))
    insert_footnote(paragraph, text, base_font_size, footnotes, footnote_per_page=5)

    # Форматирование абзаца
    apply_paragraph_format(paragraph, base_font_size, cell)

    if paragraph.runs:
        run = paragraph.runs[0]
        apply_run_format(run, base_font_size)

def add_title(doc, base_font_size, number =1 ):
    
    title_number = random.choice([True, False])
    title = doc.add_paragraph()
    run = title.add_run()

    title_text = fake.sentence(nb_words=random.randint(1, 50))
    if title_number:
        title_text = f"{random.randint(0, 20)}. {title_text}"

    run.text = title_text

    # Применение стилей заголовка
    apply_title_format(title, base_font_size, number)

def add_table(doc, base_font_size, column_width, cell=None):
    """
    Добавляет таблицу в документ
    """
    if cell is not None:
        add_table_caption(doc, base_font_size)
        table = cell.add_table(rows=random.randint(3, 5), cols=random.randint(2, 4))
    else:
        add_table_caption(doc, base_font_size)
        table = doc.add_table(rows=random.randint(3, 10), cols=random.randint(3, 7))

    remove_table_borders(table)

    # Убираем границы для таблицы
    remove_border = random.choice([True, False])
    if remove_border:
        side_borders = random.choice([True, False])
        if side_borders:
            tbl = table._tbl
            tblPr = tbl.tblPr
            borders = OxmlElement('w:tblBorders')

            for border_name in ('left', 'right', 'insideH', 'insideV'):
                border_element = OxmlElement(f'w:{border_name}')
                border_element.set(qn('w:val'), 'nil')
                borders.append(border_element)
            tblPr.append(borders)

    # Применяем стиль таблицы
    apply_style = random.choice(['Table Grid'])
    table.style = apply_style

    # Рассчитываем одинаковую ширину для всех колонок
    total_width = column_width
    column_width_table = total_width / len(table.columns)

    # Устанавливаем одинаковую ширину для каждой ячейки в колонке
    for col in table.columns:
        col.width = int(column_width_table)* 1440

    # Выравнивание таблицы
    table.alignment = random.choice([
        WD_ALIGN_PARAGRAPH.JUSTIFY,
        WD_ALIGN_PARAGRAPH.CENTER
    ])

    # Применяем цвет фона и текст
    text_color_choice = random.choice(['black', 'white'])
    if text_color_choice == 'black':
        text_rgb = RGBColor(0, 0, 0)
        row_colors = [(255, 255, 255), (220, 220, 220), (240, 248, 255)]
    else:
        text_rgb = RGBColor(255, 255, 255)
        row_colors = [(0, 0, 0), (105, 105, 105), (25, 25, 112)]

    if len(row_colors) >= 2:
        color1, color2 = random.sample(row_colors, 2)
    else:
        color1, color2 = row_colors[0], row_colors[0]

    for i, row in enumerate(table.rows):
        bg_color = color1 if i % 2 == 0 else color2
        for cell in row.cells:
            # Генерация содержимого ячейки
            text = generate_realistic_data()
            insert_footnote(cell.paragraphs[0], text, base_font_size, [], footnote_per_page=0)

            if not cell.paragraphs[0].runs:
                run = cell.paragraphs[0].add_run(cell.text)
            else:
                run = cell.paragraphs[0].runs[0]

            run.font.size = Pt(base_font_size)
            cell.paragraphs[0].alignment = random.choice([
                WD_ALIGN_PARAGRAPH.JUSTIFY,
                WD_ALIGN_PARAGRAPH.CENTER,
                WD_ALIGN_PARAGRAPH.LEFT,
                WD_ALIGN_PARAGRAPH.RIGHT
            ])

            # Цвет фона для ячейки
            tcPr = cell._element.get_or_add_tcPr()
            shd = tcPr.find(qn('w:shd'))
            if shd is None:
                shd = OxmlElement("w:shd")
                tcPr.append(shd)
            shd_color = '{:02X}{:02X}{:02X}'.format(*bg_color)
            shd.set(qn('w:fill'), shd_color)
            shd.set(qn('w:val'), 'clear')

            run.font.color.rgb = text_rgb


def add_picture(doc, base_font_size, cell=None):
    """
    Добавляет изображение в документ 
    В многоколоночных документах масштабирует изображение до ширины колонки, если оно шире.
    """
    location_signature_after = random.choices(
        [True, False],
        weights=[0.7, 0.3],
        k=1
    )[0]
    image_folder = "document_generator/Dataset_images" # директория с изображениями для всавки в документ

    if os.path.isdir(image_folder):
        images = [f for f in os.listdir(image_folder) if f.endswith(('.png', '.jpg', '.jpeg'))]
        if images:
            image_path = os.path.join(image_folder, random.choice(images))

            if cell is not None:
                paragraph = cell.add_paragraph()
            else:
                paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()

            # Вставка подписи до изображения 
            if not location_signature_after:
                add_picture_caption(doc, base_font_size, cell)

            # Определяем ширину колонки, если документ многоколоночный
            if cell is not None:
                # Получаем ширину колонки
                column_width = cell.width
                max_width = column_width
                # Устанавливаем максимальную ширину изображения, сохраняя пропорции
                try:
                    run.add_picture(image_path, width=max_width)
                except Exception as e:
                    print(f"Ошибка при вставке изображения: {e}")
            else:
                try:
                    run.add_picture(image_path, width=Inches(3))
                except Exception as e:
                    print(f"Ошибка при вставке изображения: {e}")

            # Вставка подписи после изображения 
            if location_signature_after:
                add_picture_caption(doc, base_font_size, cell)
        else:
            print("No images found in folder.")
    else:
        print("Folder is empty.")

def add_picture_caption(doc, base_font_size, cell=None):
    """
    Добавляет подпись к изображению.
    """
    img_text = random.choice(["Рис. ", "Рисунок "])
    number = str(random.randint(1, 100)) + (random.choice(["-", ""]) if img_text == "Рисунок " else "")
    caption_text = img_text + number

    if cell is not None:
        paragraph = cell.add_paragraph()
    else:
        paragraph = doc.add_paragraph()
    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]

    run = paragraph.add_run()
    run.add_text(" " + caption_text)

    run.font.size = Pt(base_font_size - random.choice([1, 2, 0, 3]))
    run.font.italic = True

def add_table_caption(doc, base_font_size, cell=None):
    """
    Добавляет подпись к таблице.
    """
    tbl_text = random.choice(["Табл. ", "Таблица "])
    number = str(random.randint(1, 100)) + (random.choice(["-", ""]) if tbl_text == "Таблица " else "")
    caption_text = tbl_text + number

    if cell is not None:
        paragraph = cell.add_paragraph()
    else:
        paragraph = doc.add_paragraph()
    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]

    run = paragraph.add_run()
    run.add_text(" " + caption_text)

    run.font.size = Pt(base_font_size - random.choice([1, 2, 0, 3]))
    run.font.italic = True

def add_numbered_list(doc, base_font_size, cell=None):
    """
    Добавляет нумерованный список в документ или в указанную ячейку таблицы.
    """
    numbering_styles = ['1.', 'I.', 'A.', 'a.', 'i.', '1)', 'I)', 'A)', 'a)', 'i)']
    style_choice = random.choice(numbering_styles) 

    line_spacing = 0.8
    left_indent = Pt(random.randint(5, 15))
    count = random.randint(2, 7)

    for i in range(1, count + 1):
        if cell is not None:
            paragraph = cell.add_paragraph()
        else:
            paragraph = doc.add_paragraph()

        if style_choice in ['1.', '1)']:
            number = f"{i}{style_choice[-1]}"
        elif style_choice in ['I.', 'I)']:
            number = f"{roman_number(i)}{style_choice[-1]}"
        elif style_choice in ['A.', 'A)']:
            number = f"{chr(64 + i)}{style_choice[-1]}"
        elif style_choice in ['a.', 'a)']:
            number = f"{chr(96 + i)}{style_choice[-1]}"
        elif style_choice in ['i.', 'i)']:
            number = f"{roman_number(i).lower()}{style_choice[-1]}"

        run = paragraph.add_run(f"{number} {fake.sentence(nb_words=random.randint(3, 10))}")

        run.font.size = Pt(base_font_size)

        paragraph.paragraph_format.line_spacing = line_spacing
        paragraph.paragraph_format.left_indent = left_indent

def add_marked_list(doc, base_font_size, cell=None):
    """
    Добавляет маркированный список в документ или в указанную ячейку таблицы.
    """
    markers = ['•', '○', '■', '–', '*', '→']
    marker = random.choice(markers)  

    line_spacing = 0.8
    left_indent = Pt(random.randint(5, 15))
    count = random.randint(2, 7)

    for _ in range(count):
        if cell is not None:
            paragraph = cell.add_paragraph()  
        else:
            paragraph = doc.add_paragraph()  
        paragraph.paragraph_format.line_spacing = line_spacing  
        paragraph.paragraph_format.left_indent = left_indent  

        text_with_marker = f"{marker} {fake.sentence(nb_words=random.randint(3, 10))}"
        run = paragraph.add_run(text_with_marker)
        run.font.size = Pt(base_font_size)

def add_header(doc, base_font_size):

    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    paragraph.text = fake.text(max_nb_chars=random.randint(5, 20))
    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]
    run = paragraph.runs[0]
    run.font.size = Pt(base_font_size - 1)
    run.font.italic = True

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    paragraph.alignment = random.choices([
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.LEFT,
        WD_ALIGN_PARAGRAPH.RIGHT
        ], weights=[0.5, 0.25, 0.25], k=1)[0]
    page_number = random.choice([True, False])
    if page_number:
        add_page_number(paragraph)
    else:
        paragraph.add_run(fake.word())

def add_footnotes_section(doc, footnotes, base_font_size):
    if not footnotes:
        return

    doc.add_paragraph("_" * 50)

    for number, text in footnotes:
        footnote_paragraph = doc.add_paragraph()
        footnote_run = footnote_paragraph.add_run(f"{to_superscript(number)}. {text}")
        footnote_run.font.size = Pt(base_font_size)
        footnote_paragraph.paragraph_format.space_before = Pt(2)

def add_formula(doc, base_font_size, cell=None):
    image_stream = generate_formula_image()

    if cell is not None:
        paragraph = cell.add_paragraph()
    else:
        paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        if cell is not None:
            column_width = cell.width
            max_width = column_width
            run.add_picture(image_stream, width=max_width)
        else:
            run.add_picture(image_stream, width=Inches(3))  
    except Exception as e:
        print(f"Ошибка при вставке формулы: {e}")
