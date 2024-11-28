from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random

def roman_number(n):
    val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
    syb = ["M", "CM", "D", "CD", "C", "XC", "L", "XL", "X", "IX", "V", "IV", "I"]
    roman_num = ''
    for i in range(len(val)):
        count = int(n / val[i])
        roman_num += syb[i] * count
        n -= val[i] * count
    return roman_num

def to_superscript(number):
    superscript_map = str.maketrans('0123456789', '⁰¹²³⁴⁵⁶⁷⁸⁹')
    return str(number).translate(superscript_map)

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def insert_footnote(paragraph, text, base_font_size, footnotes, footnote_per_page=5):
    words = text.split()
    if len(words) < 3:
        run = paragraph.add_run(text)
        run.font.size = Pt(base_font_size)
        return

    if random.random() < 0.2 and len(footnotes) < footnote_per_page:
        footnote_number = len(footnotes) + 1
        footnote_number_sup = to_superscript(footnote_number)

        insert_pos = random.randint(1, len(words)-1)
        new_text = ' '.join(words[:insert_pos]) + f' {footnote_number_sup} ' + ' '.join(words[insert_pos:])
        run = paragraph.add_run(new_text)
        run.font.size = Pt(base_font_size)

        footnote_text = text  
        footnotes.append((footnote_number, footnote_text))
    else:
        run = paragraph.add_run(' '.join(words))
        run.font.size = Pt(base_font_size)
